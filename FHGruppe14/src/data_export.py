import os
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx2pdf import convert  # Artem
from matplotlib.figure import Figure


def save_as_pdf(doc_path: str, output_dir: str) -> str:
    """
    Converts a Word document to PDF format and saves it in the specified directory.
    Returns the path to the created PDF file.
    """
    pdf_path = os.path.join(
        output_dir, os.path.splitext(os.path.basename(doc_path))[0] + ".pdf"
    )
    convert(doc_path, output_path=output_dir)
    print(f"PDF created: {pdf_path}")

    return pdf_path


def create_report_template() -> None:
    """
    Creates a standardized Word document template for weekly reports.
    The template includes placeholders for the week number, weekly report content,
    comparison with previous week, and a chart/image.
    """
    # Ensure templates directory exists
    templates_dir = os.path.join(".", "templates")
    if not os.path.exists(templates_dir):
        os.makedirs(templates_dir)

    # Create new document
    doc = Document()

    # Add title
    title = doc.add_heading(level=1)
    title_run = title.add_run("Bericht [week]")
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Define main report sections with placeholder text
    sections = [
        "Wochenbericht",
        "Vergleich zur vorherigen Woche",
    ]

    for section in sections:
        doc.add_heading(section, level=2)
        doc.add_paragraph(f"[{section.lower().replace(' ', '')}]")

    # Add placeholder for chart/visualization
    doc.add_paragraph("[image]").alignment = WD_ALIGN_PARAGRAPH.CENTER

    template_path = os.path.join(templates_dir, "report_template.docx")
    doc.save(template_path)

    print(f"Template saved at: {os.path.abspath(template_path)}")


def insert_content(
    content_dict: dict[str, str],
    image: Figure,
    idx: int,
    template_path: Optional[str] = r".\templates\report_template.docx",
) -> str:
    """
    Populates the report template with actual content and generates both DOCX and PDF versions.

    Args:
        content_dict: Dictionary containing report content, with keys matching section names
        image: Matplotlib figure to be included in the report
        idx: Unique identifier for the report
        template_path: Path to the report template file

    Returns:
        str: Path to the generated PDF file

    Raises:
        ValueError: If any placeholders remain unreplaced or if image creation fails
        FileNotFoundError: If PDF generation fails
    """
    doc = Document(template_path)

    # Track placeholder replacements to ensure completeness
    placeholders = [
        "[week]",
        "[wochenbericht]",
        "[vergleichzurvorherigenwoche]",
        "[image]",
    ]
    replaced_placeholders = set()

    # Update the week number in the report title
    for paragraph in doc.paragraphs:
        if "[week]" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace("[week]", content_dict.get("week", ""))
                replaced_placeholders.add("[week]")

    # Replace content section placeholders with actual content
    for paragraph in doc.paragraphs:
        if "[wochenbericht]" in paragraph.text:
            paragraph.text = content_dict.get("Wochenbericht", "")
            replaced_placeholders.add("[wochenbericht]")
        elif "[vergleichzurvorherigenwoche]" in paragraph.text:
            paragraph.text = content_dict.get("Vergleich zur vorherigen Woche", "")
            replaced_placeholders.add("[vergleichzurvorherigenwoche]")

        # Handle chart/visualization placement
        if "[image]" in paragraph.text:
            paragraph.text = ""
            try:
                image.savefig("temp_plot.png")
            except Exception as e:
                raise ValueError(f"Error creating image for report: {e}")

            doc.add_picture("temp_plot.png", width=Inches(6))
            os.remove("temp_plot.png")
            replaced_placeholders.add("[image]")

    # Ensure all placeholders were successfully replaced
    missing_placeholders = set(placeholders) - replaced_placeholders
    if missing_placeholders:
        raise ValueError(
            f"The following placeholders were not replaced: {missing_placeholders}"
        )

    # Generate final documents
    doc_path = f"report_final_{idx}.docx"
    doc.save(doc_path)

    output_dir = os.getcwd()
    pdf_path = save_as_pdf(doc_path, output_dir)

    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Failed to create PDF file at: {pdf_path}")

    # Clean up temporary Word document
    os.remove(doc_path)

    # Open the generated PDF
    os.startfile(pdf_path)
